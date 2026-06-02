"""Deterministic generators for ingestion test fixtures.

All fixtures are generated in memory at test time so no binary files are
committed. Everything here is clearly sample data, not real research. The
generators are deterministic: the same call produces the same bytes-shape and
the same recognizable marker text the tests assert on.

Generators:
  - born_digital_pdf: a text PDF with a known marker, for the PDF parse path.
  - scanned_pdf: an image-only PDF with no text layer, for scanned detection.
  - sample_docx: a DOCX with a paragraph and a table, for the DOCX parse path.
  - messy_text: irregular whitespace and unusual Unicode, for normalization.

reportlab (BSD) generates the PDFs and Pillow (permissive) the image. python-docx
generates the DOCX. These are dev/test-only dependencies, not shipped.
"""

from __future__ import annotations

import io

# Recognizable sample markers the tests assert on. Clearly labeled as sample.
PDF_MARKER = "SAMPLE born-digital PDF for the IB Desk ingestion test."
DOCX_PARAGRAPH_MARKER = "SAMPLE docx paragraph for IB Desk ingestion test."
DOCX_TABLE_MARKER = "SampleTableCellValue"


def born_digital_pdf() -> bytes:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=LETTER)
    pdf.setFont("Helvetica", 12)
    pdf.drawString(72, 720, PDF_MARKER)
    pdf.drawString(72, 700, "This is sample data, not real research.")
    pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def scanned_pdf() -> bytes:
    """An image-only PDF: a blank raster image on the page, no text layer."""
    from PIL import Image
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen import canvas

    image = Image.new("RGB", (600, 800), color="white")
    image_buffer = io.BytesIO()
    image.save(image_buffer, format="PNG")
    image_buffer.seek(0)

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=LETTER)
    pdf.drawImage(ImageReader(image_buffer), 0, 0, width=LETTER[0], height=LETTER[1])
    pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def sample_docx() -> bytes:
    from docx import Document as DocxDocument

    document = DocxDocument()
    document.add_paragraph(DOCX_PARAGRAPH_MARKER)
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = DOCX_TABLE_MARKER
    table.rows[0].cells[1].text = "second cell"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def messy_text() -> str:
    return "  Hello  world \r\n\r\n\r\n  Second\tparagraph\x00 here  \r\n   "
