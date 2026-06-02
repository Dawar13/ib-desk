"""Deterministic parsing of PDF, DOCX, and pasted text into page segments.

Parsing extracts text only. It does not normalize (normalization.assemble_raw_text
does that) and it does not apply policy: it reports the page count and the count
of extractable non-whitespace characters so the caller can apply the configurable
scanned threshold. The same input always produces the same output, with no
randomness.

The scanned or empty case is reported through the metrics on ParseResult rather
than by silently returning empty text, so the caller can reject it honestly. OCR
is deferred (see BUILD_PLAN.md and the Phase 1 spec).
"""

from __future__ import annotations

import io
from dataclasses import dataclass

import pdfplumber
from docx import Document as DocxDocument


@dataclass(frozen=True)
class ParseResult:
    # One entry per PDF page, or a single entry for DOCX and pasted text.
    segments: list[str]
    page_count: int
    # Total count of non-whitespace characters extracted, used for scanned detection.
    extractable_chars: int


def _count_nonspace(segments: list[str]) -> int:
    return sum(len("".join(segment.split())) for segment in segments)


def parse_pdf(data: bytes) -> ParseResult:
    segments: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            segments.append(page.extract_text() or "")
    return ParseResult(
        segments=segments,
        page_count=len(segments),
        extractable_chars=_count_nonspace(segments),
    )


def parse_docx(data: bytes) -> ParseResult:
    document = DocxDocument(io.BytesIO(data))
    parts: list[str] = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    text = "\n".join(parts)
    return ParseResult(
        segments=[text],
        page_count=1,
        extractable_chars=_count_nonspace([text]),
    )


def parse_text(text: str) -> ParseResult:
    return ParseResult(
        segments=[text],
        page_count=1,
        extractable_chars=_count_nonspace([text]),
    )


def looks_scanned(result: ParseResult, min_chars_per_page: int) -> bool:
    """Return True when the document yields effectively no extractable text.

    The threshold is a tunable heuristic (the average extractable characters per
    page). A fully scanned, image-only PDF yields zero extractable characters and
    is always treated as scanned.
    """
    if result.extractable_chars == 0:
        return True
    if result.page_count <= 0:
        return True
    return (result.extractable_chars / result.page_count) < min_chars_per_page
